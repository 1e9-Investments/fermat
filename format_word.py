#!/usr/bin/env python3
"""
Fermat Word Formatter - Applies 1e9-Investments document standards
Full spec: margins, fonts, headings, bullets, tables, spacing, style fixes
"""

import sys
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# === COLOR PALETTE (Greenish - Wall St Standard) ===
COLORS = {
    'h1': RGBColor(0x1A, 0x3C, 0x2B),           # Deep forest green
    'h2': RGBColor(0x2D, 0x6A, 0x4F),           # Mid forest green
    'h3': RGBColor(0x40, 0x91, 0x6C),           # Medium green
    'sublabel': RGBColor(0x52, 0x79, 0x6F),     # Muted sage (italic sub-labels)
    'source': RGBColor(0x59, 0x59, 0x59),       # Gray for sources
    'table_header': RGBColor(0x1F, 0x4E, 0x36), # Dark forest green header fill
    'key_row': RGBColor(0xC2, 0xD6, 0x9B),      # Sage highlight for key metrics
    'alt_row': RGBColor(0xF5, 0xF5, 0xF5),      # Light gray alternating
    'border': RGBColor(0xF2, 0xF2, 0xF2),       # Near-invisible border
}

# Key metric row keywords
KEY_METRICS = ['revenue', 'pat', 'profit after tax', 'net income', 'total debt', 'total assets',
               'ebitda', 'gross profit', 'operating income', 'net profit']

# === FILLER PHRASES TO REMOVE ===
FILLER_PATTERNS = [
    r'\bIt is worth noting that\b',
    r'\bIt should be noted that\b',
    r'\bIt is important to note that\b',
    r'\bAs mentioned earlier,?\s*',
    r'\bAs previously stated,?\s*',
    r'\bIn order to\b',
    r'\bDue to the fact that\b',
    r'\bAt the present time\b',
    r'\bAt this point in time\b',
    r'\bIn the event that\b',
    r'\bFor the purpose of\b',
    r'\bWith regard to\b',
    r'\bWith respect to\b',
    r'\bIn terms of\b',
    r'\bOn the other hand,?\s*',
    r'\bHaving said that,?\s*',
    r'\bThat being said,?\s*',
    r'\bAll things considered,?\s*',
    r'\bBy and large,?\s*',
    r'\bFor all intents and purposes,?\s*',
    r'\bBasically,?\s*',
    r'\bEssentially,?\s*',
    r'\bActually,?\s*',
    r'\bObviously,?\s*',
    r'\bClearly,?\s*',
    r'\bGenerally speaking,?\s*',
    r'\bIt goes without saying that\b',
    r'\bNeedless to say,?\s*',
    r'\bAs a matter of fact,?\s*',
    r'\bIn my opinion,?\s*',
    r'\bI think that\b',
    r'\bI believe that\b',
]

# === HEDGING TO STRENGTHEN ===
HEDGE_REPLACEMENTS = {
    r'\bperhaps\b': '',
    r'\bmaybe\b': '',
    r'\bpossibly\b': '',
    r'\bmight be\b': 'is',
    r'\bcould be\b': 'is',
    r'\bseems to be\b': 'is',
    r'\bappears to be\b': 'is',
    r'\btends to\b': '',
    r'\bin some cases\b': '',
    r'\bto some extent\b': '',
    r'\brelatively\b': '',
    r'\bsomewhat\b': '',
    r'\bslightly\b': '',
    r'\bfairly\b': '',
    r'\brather\b': '',
    r'\bquite\b': '',
    r'\ba bit\b': '',
    r'\bkind of\b': '',
    r'\bsort of\b': '',
}


def set_cell_shading(cell, color_hex):
    """Set cell background color"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_borders(cell, color_hex='F2F2F2', width='4'):
    """Set cell borders - 0.5pt single, light gray"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), width)  # 4 = 0.5pt
        border.set(qn('w:color'), color_hex)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def fix_em_dashes(text):
    """Replace em dashes and en dashes with hyphens"""
    if not text:
        return text
    return text.replace('—', '-').replace('–', '-')


def fix_financial_notation(text):
    """Standardize financial notation: US$Xmm (c.VNDXbn), bn/tn, Xx ratios"""
    if not text:
        return text

    # Fix billion/million notation - lowercase bn/mm/tn
    text = re.sub(r'\$(\d+(?:\.\d+)?)\s*[Bb](?:illion)?(?!\w)', r'US$\1bn', text)
    text = re.sub(r'\$(\d+(?:\.\d+)?)\s*[Mm](?:illion)?(?!\w)', r'US$\1mm', text)
    text = re.sub(r'\$(\d+(?:\.\d+)?)\s*[Tt](?:rillion)?(?!\w)', r'US$\1tn', text)

    # Fix VND notation
    text = re.sub(r'VND\s*(\d+(?:\.\d+)?)\s*[Bb](?:illion)?(?!\w)', r'VND\1bn', text)
    text = re.sub(r'VND\s*(\d+(?:\.\d+)?)\s*[Mm](?:illion)?(?!\w)', r'VND\1mm', text)
    text = re.sub(r'VND\s*(\d+(?:\.\d+)?)\s*[Tt](?:rillion)?(?!\w)', r'VND\1tn', text)

    # Standardize standalone B/M/T to bn/mm/tn
    text = re.sub(r'(\d)B\b', r'\1bn', text)
    text = re.sub(r'(\d)M\b', r'\1mm', text)
    text = re.sub(r'(\d)T\b', r'\1tn', text)

    # Fix ratio notation (e.g., 2x, 3.5x)
    text = re.sub(r'(\d+(?:\.\d+)?)\s*[Xx]\b', r'\1x', text)

    # Fix year labels (2025A, 2026F, 2026E)
    text = re.sub(r'\b(20\d{2})\s*[Aa](?:ctual)?(?!\w)', r'\1A', text)
    text = re.sub(r'\b(20\d{2})\s*[Ff](?:orecast)?(?!\w)', r'\1F', text)
    text = re.sub(r'\b(20\d{2})\s*[Ee](?:stimate)?(?!\w)', r'\1F', text)

    # Fix growth notation (+X% YoY)
    text = re.sub(r'(\d+(?:\.\d+)?%)\s*[Yy]o[Yy]', r'\1 YoY', text)

    return text


def remove_filler(text):
    """Remove filler phrases"""
    if not text:
        return text
    for pattern in FILLER_PATTERNS:
        text = re.sub(pattern, '', text, flags=re.IGNORECASE)
    # Clean up double spaces
    text = re.sub(r'\s+', ' ', text).strip()
    # Fix sentence starts after removal
    text = re.sub(r'\.\s+([a-z])', lambda m: '. ' + m.group(1).upper(), text)
    return text


def strengthen_hedging(text):
    """Remove or strengthen hedging language"""
    if not text:
        return text
    for pattern, replacement in HEDGE_REPLACEMENTS.items():
        text = re.sub(pattern, replacement, text, flags=re.IGNORECASE)
    # Clean up
    text = re.sub(r'\s+', ' ', text).strip()
    return text


def process_text_content(text):
    """Apply all content/style fixes"""
    if not text:
        return text
    text = fix_em_dashes(text)
    text = fix_financial_notation(text)
    text = remove_filler(text)
    text = strengthen_hedging(text)
    return text


def apply_font_to_run(run, font_name='Aptos Narrow', size_pt=11):
    """Apply font settings to a run - Aptos Narrow primary"""
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    # Set fallback fonts
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:cs'), 'Calibri')  # Fallback for complex scripts


def set_paragraph_spacing(para, before_pt=6, after_pt=6, line_spacing=1.0):
    """Set paragraph spacing - 6pt before/after, 1.0x line spacing"""
    pf = para.paragraph_format
    pf.space_before = Pt(before_pt)
    pf.space_after = Pt(after_pt)
    pf.line_spacing = line_spacing
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE


def set_keep_with_next(para, keep=True):
    """Set keepNext to prevent orphaned headers"""
    pPr = para._element.get_or_add_pPr()
    keepNext = OxmlElement('w:keepNext')
    if keep:
        keepNext.set(qn('w:val'), '1')
    pPr.append(keepNext)


def remove_blank_paragraphs(doc):
    """Remove blank paragraphs used as spacers"""
    paragraphs_to_check = list(doc.paragraphs)
    for para in paragraphs_to_check:
        if not para.text.strip() and not para.runs:
            # Empty paragraph - check if it's being used as spacer
            p = para._element
            p.getparent().remove(p)


def detect_heading_level(para):
    """Detect heading level from style or formatting"""
    style_name = para.style.name.lower() if para.style else ''

    if 'heading 1' in style_name or style_name == 'h1':
        return 1
    elif 'heading 2' in style_name or style_name == 'h2':
        return 2
    elif 'heading 3' in style_name or style_name == 'h3':
        return 3

    # Check if it looks like a heading (short, bold)
    text = para.text.strip()
    if len(text) < 100 and len(text) > 0:
        all_bold = all(run.bold for run in para.runs if run.text.strip())
        if all_bold and not text.startswith(('•', '-', '○', '*', '►')):
            # Check for numbering pattern
            if re.match(r'^\d+\.', text):
                return 1
            elif re.match(r'^\d+\.\d+', text):
                return 2
            elif re.match(r'^\d+\.\d+\.\d+', text):
                return 3
            return 2  # Default to H2 for bold short text

    return 0


def is_sublabel(para):
    """Detect if paragraph is a sub-label (table title, section callout)"""
    text = para.text.strip()
    # Short, italic, often before a table
    if len(text) < 80:
        all_italic = all(run.italic for run in para.runs if run.text.strip())
        if all_italic and not any(run.bold for run in para.runs if run.text.strip()):
            return True
    # Common sub-label patterns
    if re.match(r'^(Table|Figure|Exhibit|Chart)\s*\d*[:\.]', text, re.IGNORECASE):
        return True
    return False


def is_source_line(para):
    """Detect Sources: line"""
    text = para.text.strip().lower()
    return text.startswith('source')


def is_bullet(para):
    """Detect if paragraph is a bullet"""
    text = para.text.strip()
    return text.startswith(('•', '-', '○', '*', '►', '–'))


def get_bullet_level(para):
    """Get bullet indentation level (0 or 1)"""
    text = para.text.strip()
    if text.startswith('○'):
        return 1
    # Check left indent
    if para.paragraph_format.left_indent:
        indent_inches = para.paragraph_format.left_indent.inches if para.paragraph_format.left_indent else 0
        if indent_inches >= 0.9:
            return 1
    return 0


def format_heading(para, level):
    """Format heading according to level with greenish palette"""
    set_paragraph_spacing(para, before_pt=12, after_pt=6)
    set_keep_with_next(para, True)

    for run in para.runs:
        apply_font_to_run(run, size_pt=11)
        run.text = process_text_content(run.text)

        if level == 1:
            run.bold = True
            run.underline = True
            run.italic = False
            run.font.color.rgb = COLORS['h1']
        elif level == 2:
            run.bold = True
            run.underline = False
            run.italic = False
            run.font.color.rgb = COLORS['h2']
        elif level == 3:
            run.bold = True
            run.italic = True
            run.underline = False
            run.font.color.rgb = COLORS['h3']


def format_sublabel(para):
    """Format sub-label (table title, section callout) - italic, muted sage"""
    set_paragraph_spacing(para, before_pt=6, after_pt=0)  # Glued to table below
    set_keep_with_next(para, True)

    for run in para.runs:
        apply_font_to_run(run, size_pt=11)
        run.text = process_text_content(run.text)
        run.italic = True
        run.bold = False
        run.underline = False
        run.font.color.rgb = COLORS['sublabel']


def format_source_line(para):
    """Format Sources: line - italic, 9pt, gray"""
    set_paragraph_spacing(para, before_pt=6, after_pt=6)

    for run in para.runs:
        apply_font_to_run(run, size_pt=9)
        run.text = process_text_content(run.text)
        run.italic = True
        run.bold = False
        run.font.color.rgb = COLORS['source']


def format_bullet(para, level=0):
    """Format bullet - 2 levels max, specific indents"""
    set_paragraph_spacing(para, before_pt=6, after_pt=6)

    pf = para.paragraph_format
    if level == 0:
        pf.left_indent = Inches(0.5)
        pf.first_line_indent = Inches(-0.25)  # Hanging indent
    else:
        pf.left_indent = Inches(1.0)
        pf.first_line_indent = Inches(-0.25)

    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for run in para.runs:
        apply_font_to_run(run, size_pt=11)
        run.text = process_text_content(run.text)


def format_body(para):
    """Format regular body paragraph"""
    set_paragraph_spacing(para, before_pt=6, after_pt=6)
    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.left_indent = Inches(0)
    para.paragraph_format.first_line_indent = Inches(0)

    for run in para.runs:
        apply_font_to_run(run, size_pt=11)
        run.text = process_text_content(run.text)


def is_key_metric_row(row_text):
    """Check if row contains key metrics"""
    text_lower = row_text.lower()
    return any(metric in text_lower for metric in KEY_METRICS)


def format_table(table):
    """Format table according to spec"""
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    num_rows = len(table.rows)

    for row_idx, row in enumerate(table.rows):
        row_text = ' '.join(cell.text for cell in row.cells)
        is_header = row_idx == 0
        is_key_metric = is_key_metric_row(row_text) if not is_header else False
        is_alt_row = row_idx % 2 == 0 and not is_header

        for cell_idx, cell in enumerate(row.cells):
            # Set borders - 0.5pt light gray
            set_cell_borders(cell, 'F2F2F2', '4')

            # Format cell paragraphs
            for para in cell.paragraphs:
                # Cell spacing: 0pt line spacing for maximum density
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(0)
                para.paragraph_format.line_spacing = 1.0

                for run in para.runs:
                    apply_font_to_run(run, size_pt=11)
                    run.text = fix_em_dashes(fix_financial_notation(run.text)) if run.text else ''

                    if is_header:
                        run.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
                    else:
                        run.font.color.rgb = RGBColor(0, 0, 0)

            # Cell shading
            if is_header:
                set_cell_shading(cell, '1F4E36')  # Dark forest green
            elif is_key_metric:
                set_cell_shading(cell, 'C2D69B')  # Sage highlight
            elif is_alt_row:
                set_cell_shading(cell, 'F5F5F5')  # Light gray
            else:
                set_cell_shading(cell, 'FFFFFF')  # White

            # Alignment: right for numbers, left for labels
            cell_text = cell.text.strip()
            is_numeric = bool(re.match(r'^[\d,.\-\(\)%$VNDUSmmbntk\s~c\.]+$', cell_text)) if cell_text else False

            for para in cell.paragraphs:
                if is_numeric and cell_idx > 0:  # First column usually labels
                    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT


def set_margins(doc):
    """Set page to US Letter with 0.5in margins"""
    for section in doc.sections:
        # US Letter size
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        # 0.5in margins
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.gutter = Inches(0)


def format_document(filepath):
    """Main formatting function"""
    print(f"Loading: {filepath}")
    doc = Document(filepath)

    # Set margins
    set_margins(doc)

    # Process paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()

        if not text:
            continue

        # Check for Sources line first
        if is_source_line(para):
            format_source_line(para)
            continue

        # Check for sub-label
        if is_sublabel(para):
            format_sublabel(para)
            continue

        # Check for heading
        heading_level = detect_heading_level(para)
        if heading_level > 0:
            format_heading(para, heading_level)
            continue

        # Check for bullet
        if is_bullet(para):
            level = get_bullet_level(para)
            format_bullet(para, level)
            continue

        # Regular body paragraph
        format_body(para)

    # Process tables
    for table in doc.tables:
        format_table(table)

    # Remove blank paragraphs used as spacers (do this last)
    # Note: commenting out for now to avoid removing intentional spacing
    # remove_blank_paragraphs(doc)

    # Save
    doc.save(filepath)
    print(f"Saved: {filepath}")


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("Usage: format_word.py <filepath>")
        sys.exit(1)

    filepath = sys.argv[1]
    if not Path(filepath).exists():
        print(f"Error: File not found: {filepath}")
        sys.exit(1)

    format_document(filepath)
